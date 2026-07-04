from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('丝印外发加工协议')
run.bold = True
run.font.size = Pt(18)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 协议抬头信息
def add_info(doc, label, val=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label)
    run.font.size = Pt(11)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if val:
        run2 = p.add_run(val)
        run2.font.size = Pt(11)

add_info(doc, '协议编号：_______________')
add_info(doc, '签订地点：_______________')
add_info(doc, '签订日期：2026年7月1日')
add_info(doc, '')

p = doc.add_paragraph()
run = p.add_run('甲方（委托方）：_______________')
run.font.size = Pt(11)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
run = p.add_run('乙方（加工方）：_______________')
run.font.size = Pt(11)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
run = p.add_run('甲、乙双方根据《中华人民共和国民法典》及相关法律法规，本着平等互利、诚实信用的原则，就甲方委托乙方进行丝印加工事宜，经友好协商达成如下协议：')
run.font.size = Pt(11)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

sections = [
    ("第一条  合作期限", [
        "本协议合作期限为壹年，自2026年7月1日起至2027年6月30日止。协议期满前一个月，双方协商是否续签。"
    ]),
    ("第二条  加工内容及要求", [
        "1. 甲方委托乙方进行丝网印刷加工，具体加工内容以甲方出具的《加工订单》或双方微信对接群内确认的信息为准。",
        "2. 加工质量、工艺标准及验收要求以甲方提供的样品或技术文件为准，乙方须严格按照要求生产。"
    ]),
    ("第三条  网框提供", [
        "1. 合作初期：印刷所需网框由甲方提供。",
        "2. 后期变更：后续若乙方自行制作网框，经甲方确认合格后，可交由乙方制作。",
        "3. 网框所有权归甲方所有，乙方应妥善保管，如有损坏或丢失，乙方按成本价赔偿。"
    ]),
    ("第四条  加工费用及计价方式", [
        "1. 单价：人民币 0.4 元/印（大写：肆角整/印）。",
        "2. 计价规则：",
        "   (1) 不论版面大小、数量，每印刷一次（一印）按0.4元计费。",
        "   (2) 多色印刷：每增加一色按一印计算（如2色=2印）。",
        "   (3) 里外印刷：正反两面印刷按2印计算。",
        "   (4) 多版印刷：超大内容需分版印刷的，每增加一个版按一印计算（如2个版=2印）。",
        "3. 以上单价为含税价。"
    ]),
    ("第五条  油墨供应", [
        "1. 常规油墨：由乙方负责提供，费用已包含在加工费中。",
        "2. 特殊油墨：由甲方自行提供，乙方负责按甲方要求使用。"
    ]),
    ("第六条  货物交接与包装", [
        "1. 提货：乙方自行前往甲方指定地点提取加工原材料/半成品，相关运输费用由乙方承担。",
        "2. 送货：加工完成后，乙方负责将成品送回甲方指定地点，相关运输费用由乙方承担。",
        "3. 包装：加工完成后的打包工作由乙方自行负责，打包所需耗材由乙方自行解决，包装须符合运输和存储要求。"
    ]),
    ("第七条  交期", [
        "1. 常规交期：自乙方提货之日起3个自然日内完成加工并交付。",
        "2. 加急订单：由双方在微信对接群内另行协商确定交期。",
        "3. 所有订单的交期确认以微信对接群（群名称：_______________）内的沟通记录为准。"
    ]),
    ("第八条  结算方式", [
        "1. 对账周期：每月1日至当月最后一日为一个结算周期。",
        "2. 对账时间：次月10日前，双方完成上月加工数量及金额的对账确认。",
        "3. 付款时间：次月20日前，甲方支付上月对账确认的全部金额。",
        "4. 乙方应在对账后向甲方开具合法有效的增值税专用发票，甲方在收到发票后按约付款。"
    ]),
    ("第九条  业务对接", [
        "1. 甲方对接人：姚斌（姚经理）",
        "2. 对接方式：双方通过微信对接群进行日常订单沟通、交期确认、异常反馈等业务往来。",
        "3. 双方对接人如有变更，应提前三个工作日书面通知对方。"
    ]),
    ("第十条  质量标准与验收", [
        "1. 乙方应确保加工产品符合甲方确认的质量标准，包括但不限于：印刷位置准确、颜色一致、无脏污、无漏印、附着力达标等。",
        "2. 甲方有权对乙方加工的产品进行抽检或全检，不合格产品乙方应无条件返工或重印，因此造成的交期延误由乙方承担。"
    ]),
    ("第十一条  印错/质量问题的责任承担", [
        "1. 已发货给客户后发现错误：乙方须按该批次产品的终端销售价格全额赔偿甲方，并承担因此产生的往返快递费。",
        "2. 未发货发现错误（返工可解决的）：乙方承担材料成本价。",
        "3. 大批次批量错误（未发货）：按该批次材料成本价的1.3倍（即生产成本）赔偿。",
        "4. 以上赔偿不影响甲方依据本协议其他条款追究乙方违约责任的权利。"
    ]),
    ("第十二条  版框（网框）的损坏责任", [
        "1. 首次提供：甲方仅向乙方提供一次印刷版框，后续因正常使用寿命耗尽需更换的，由甲方重新购买。",
        "2. 操作不当损坏：在生产过程中因乙方操作不当导致版框损坏的，由乙方负责赔偿。",
        "3. 保存不当损坏：因乙方保管、存放不当导致版框损坏或报废的，由乙方负责赔偿。",
        "4. 乙方应赔偿的版框金额按版框实际制作成本价计算。"
    ]),
    ("第十三条  违约责任", [
        "1. 乙方逾期交货的，每逾期一天，按该批次加工费的3%向甲方支付违约金。",
        "2. 甲方逾期付款的，每逾期一天，按应付未付金额的3‰向乙方支付违约金。",
        "3. 任何一方未经对方同意擅自解除协议，应向对方支付违约金人民币_______________元。"
    ]),
    ("第十四条  保密条款", [
        "双方对在合作过程中知悉的对方商业秘密（包括但不限于客户信息、产品工艺、价格信息等）负有保密义务，未经对方书面同意不得向第三方披露。违反本条款的，违约方应赔偿由此给对方造成的全部损失。"
    ]),
    ("第十五条  争议解决", [
        "本协议履行过程中发生争议，双方应友好协商解决；协商不成的，向甲方所在地人民法院提起诉讼。"
    ]),
    ("第十六条  其他", [
        "1. 本协议一式两份，甲乙双方各执一份，具有同等法律效力。",
        "2. 本协议自双方签字（或盖章）之日起生效。",
        "3. 未尽事宜，双方可另行签订补充协议。"
    ]),
]

for title_text, items in sections:
    # 标题
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title_text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 内容
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 分隔线
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)

# 签名区
p = doc.add_paragraph()
run = p.add_run('甲方（盖章）：_______________')
run.font.size = Pt(11)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
run = p.add_run('授权代表签字：_______________')
run.font.size = Pt(11)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
run = p.add_run('联系电话：_______________')
run.font.size = Pt(11)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
run = p.add_run('日期：______年______月______日')
run.font.size = Pt(11)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)

p = doc.add_paragraph()
run = p.add_run('乙方（盖章）：_______________')
run.font.size = Pt(11)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
run = p.add_run('授权代表签字：_______________')
run.font.size = Pt(11)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
run = p.add_run('联系电话：_______________')
run.font.size = Pt(11)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
run = p.add_run('日期：______年______月______日')
run.font.size = Pt(11)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 保存
output_path = r'D:\Desktop\sanyang-system\丝印外发加工协议.docx'
doc.save(output_path)
print(f"已保存: {output_path}")
print(f"文件大小: {os.path.getsize(output_path) / 1024:.1f} KB")
